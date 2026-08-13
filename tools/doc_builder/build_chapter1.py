"""Generate Chapter 1 Introduction as Word document."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "reports" / "EthiMatch_Chapter1_Introduction.docx"

NAVY = RGBColor(0x0B, 0x24, 0x47)


def add_heading(doc, text, level=1):
    h = doc.add_heading(level=level)
    r = h.add_run(text)
    r.font.color.rgb = NAVY
    r.font.name = "Calibri"
    r.font.bold = True
    if level == 1:
        r.font.size = Pt(16)
    else:
        r.font.size = Pt(13)


def add_para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.font.color.rgb = NAVY


def main():
    doc = Document()
    for s in doc.sections:
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

    add_heading(doc, "Chapter 1: Introduction", 1)

CHAPTER_1_SECTIONS = [
        ("1.1 Background and Context", (
            "Clinical trials are essential for evaluating new cancer treatments, yet recruiting "
            "eligible patients remains one of the most persistent operational bottlenecks in "
            "oncology research. Coordinators and oncologists must manually review lengthy "
            "electronic health records (EHRs), interpret unstructured clinical notes, and "
            "compare each patient's profile against complex inclusion and exclusion criteria "
            "defined in trial protocols. This process is time-consuming, costly, and vulnerable "
            "to human inconsistency. Carlisle et al. (2015) analysed 2,579 phase 2 and 3 trials "
            "closed in 2011 and found that 481 (19%) either terminated because of failed accrual "
            "or completed with less than 85% of expected enrolment, seriously compromising "
            "statistical power. Automated patient–trial matching therefore represents a "
            "significant opportunity to widen the recruitment funnel while reducing the "
            "administrative burden on clinical research teams."
            "\n\n"
            "Recent advances in biomedical natural language processing (NLP) have made it "
            "feasible to extract clinical entities—such as diagnoses, disease stage, "
            "biomarkers, and prior therapies—directly from free-text notes (Lee et al., 2020). "
            "However, purely neural approaches carry inherent risks in a safety-critical "
            "domain. A model may misread negated statements (for example, interpreting "
            "\"no history of diabetes\" as evidence of diabetes), omit required fields, or "
            "produce confident but incorrect eligibility predictions. Zitianellis (2025) "
            "emphasises that clinical adoption of artificial intelligence (AI) in trial "
            "pre-screening depends not only on accuracy but also on transparency, "
            "interpretability, and conservative handling of missing information. These "
            "concerns motivate a design that separates text understanding from eligibility "
            "decision-making."
        )),
        ("1.2 Problem Statement", (
            "The core problem addressed by this project is whether automated clinical-trial "
            "matching can be made both efficient and safe. Efficiency requires reading "
            "unstructured notes at scale; safety requires that every protocol rule be applied "
            "deterministically, that missing data never be fabricated, and that every decision "
            "be explainable to a human reviewer. Purely neural systems optimise pattern "
            "recognition but cannot guarantee protocol compliance. Purely rule-based systems "
            "can enforce criteria strictly but cannot interpret narrative clinical text. "
            "Neither approach alone adequately serves oncologists and research coordinators "
            "who need rapid, trustworthy decision support before contacting patients."
            "\n\n"
            "Loaiza-Bonilla et al. (2026) demonstrated in a prospective evaluation of 3,804 "
            "cancer patients that neuro-symbolic architectures—combining neural extraction "
            "with strict symbolic rule logic—materially improve matching safety compared with "
            "generative AI alone. Yet few open, reproducible academic implementations exist "
            "that operationalise this concept with explicit handling of inconclusive cases, "
            "dual synthetic and real-world evaluation cohorts, and statistically rigorous "
            "comparison against a neural baseline. This project seeks to fill that gap through "
            "the design, implementation, and evaluation of EthiMatch."
        )),
        ("1.3 Research Question, Aim and Objectives", (
            "The research question guiding this project is:"
            "\n\n"
            "Can a neuro-symbolic AI architecture combining a biomedical text-reading model "
            "with a strict, deterministic rule engine improve patient safety and "
            "explainability in clinical-trial matching compared with a purely neural AI "
            "approach?"
            "\n\n"
            "The aim of the project is to design, implement, and empirically evaluate "
            "EthiMatch—a neuro-symbolic clinical-trial matching system that reads "
            "biomedical notes, extracts structured eligibility features, applies "
            "protocol-grounded rules, and presents explainable audit reports to clinical "
            "users through an interactive dashboard."
            "\n\n"
            "To achieve this aim, the following objectives were defined:"
            "\n\n"
            "1. To formulate a five-stage neuro-symbolic pipeline architecture that separates "
            "neural entity extraction from symbolic eligibility validation and explainable "
            "reporting."
            "\n"
            "2. To implement a unified data-ingestion layer supporting Synthea synthetic "
            "cohorts and the publicly available MIMIC-IV Demo dataset, normalising records "
            "into a common patient-profile contract."
            "\n"
            "3. To integrate a biomedical named-entity recognition (NER) model "
            "(d4data/biomedical-ner-all) with a deterministic symbolic validator driven by "
            "JSON trial protocols, including explicit INCONCLUSIVE verdict semantics for "
            "missing data."
            "\n"
            "4. To develop a Streamlit clinician dashboard providing audit trails, "
            "explainability visualisations, and cohort-screening functionality."
            "\n"
            "5. To benchmark the neuro-symbolic pipeline against a pure-neural baseline "
            "using precision, recall, F1-score, false positive rate (FPR), and McNemar's "
            "paired significance test across multiple datasets."
            "\n"
            "6. To critically reflect on the strengths, limitations, and ethical implications "
            "of the proposed approach within the context of MSc-level clinical AI research."
        )),
        ("1.4 Significance and Contribution", (
            "This project is significant for three groups of stakeholders. Primary users—"
            "oncologists and clinical research coordinators—benefit from reduced manual "
            "chart-review time and protocol-grounded evidence before patient contact. "
            "Secondary stakeholders, including hospital research departments and clinical "
            "informatics teams, gain a prototype decision-support tool with a clear audit "
            "trail suitable for demonstration and further integration. From an academic "
            "perspective, the project contributes an open, reproducible neuro-symbolic "
            "implementation with dual-dataset evaluation, hash-validated caching for "
            "responsive interaction, and paired statistical testing that isolates the "
            "symbolic layer's contribution while holding the neural extractor constant."
        )),
        ("1.5 Scope and Delimitations", (
            "The project is deliberately scoped as a research prototype rather than a "
            "deployable clinical product. In scope are: the EthiMatch software artefact "
            "(Python); JSON-defined oncology trial protocols; evaluation on Synthea and "
            "MIMIC-IV Demo data; and a comparative benchmarking harness producing "
            "quantitative safety and accuracy metrics. Out of scope are: live hospital EHR "
            "integration; regulatory medical-device approval; real patient enrolment; and "
            "credentialled access to the full MIMIC-IV corpus (although the architecture "
            "supports extension should access be granted). The clinical vocabulary is "
            "limited to a defined oncology subset (for example, NSCLC, SCLC, breast cancer) "
            "with corresponding disease stages and biomarkers, which is appropriate for a "
            "focused MSc prototype but not representative of all therapeutic areas."
        )),
        ("1.6 Structure of the Report", (
            "The remainder of this report is organised as follows. Chapter 2 presents a "
            "critical literature review and theoretical framework underpinning neuro-symbolic "
            "clinical-trial matching. Chapter 3 describes the research design and "
            "methodology, including datasets, evaluation metrics, and statistical procedures. "
            "Chapter 4 details the design, implementation, and testing of the EthiMatch "
            "artefact. Chapter 5 discusses project management, planning, and risk mitigation. "
            "Chapter 6 presents evaluation results and critical discussion. Chapter 7 "
            "addresses legal, ethical, and social considerations. Chapter 8 concludes with "
            "key findings, limitations, and recommendations for future work."
        )),
    ]

    for title, body in CHAPTER_1_SECTIONS:
        add_heading(doc, title, 2)
        for para in body.split("\n\n"):
            add_para(doc, para.strip())

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    word_count = sum(len(p.split()) for _, b in CHAPTER_1_SECTIONS for p in b.split("\n\n"))
    print(f"Saved: {OUT}")
    print(f"Approximate word count: {word_count}")


if __name__ == "__main__":
    main()
