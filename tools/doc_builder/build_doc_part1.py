"""Part 1 of the design-document builder.

Defines styling helpers and the cover/TOC/intro sections.
Imported by build_doc.py which orchestrates the full build.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches, Cm

PROJECT_ROOT = Path(__file__).resolve().parents[2]
ASSETS = PROJECT_ROOT / "docs" / "figures"

NAVY    = RGBColor(0x0B, 0x24, 0x47)
TEAL    = RGBColor(0x19, 0xA7, 0xCE)
GOLD    = RGBColor(0xF0, 0xA0, 0x4B)
GREEN   = RGBColor(0x3A, 0x7D, 0x44)
RED     = RGBColor(0xC5, 0x30, 0x30)
GREY    = RGBColor(0x4A, 0x55, 0x68)
LIGHT   = RGBColor(0xF5, 0xF7, 0xFA)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
PURPLE  = RGBColor(0xA0, 0x84, 0xDC)


def hex_for(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def shade_cell(cell, hex_fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def set_cell_borders(cell, color: str = "B0B7C3") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:color"), color)
        borders.append(b)
    tcPr.append(borders)


def add_horizontal_rule(doc, color: RGBColor = NAVY) -> None:
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), hex_for(color))
    pbdr.append(bottom)
    p_pr.append(pbdr)


def style_heading(doc, text: str, level: int = 1, color: RGBColor = NAVY) -> None:
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(22)
    elif level == 2:
        run.font.size = Pt(16)
    else:
        run.font.size = Pt(13)


def body_para(doc, text: str, *, size: int = 11, color: RGBColor = None,
              bold: bool = False, italic: bool = False,
              align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after: int = 6) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic


def callout(doc, title: str, body: str, *, fill_hex: str = "E3F2FD",
            border_hex: str = "0B2447", title_color: RGBColor = NAVY) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    shade_cell(cell, fill_hex)
    set_cell_borders(cell, color=border_hex)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = title_color
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    r2 = p2.add_run(body)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = NAVY
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def bullet(doc, text: str, *, size: int = 11, indent: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.left_indent = Cm(0.6 + 0.5 * indent)
    run = p.runs[0] if p.runs else p.add_run("")
    p.runs[0].text = text
    p.runs[0].font.size = Pt(size)
    p.runs[0].font.name = "Calibri"
    p.runs[0].font.color.rgb = NAVY


def make_table(doc, headers: list[str], rows: list[list[str]],
               *, header_fill: str = "0B2447",
               zebra: tuple[str, str] = ("F5F7FA", "FFFFFF"),
               header_text: RGBColor = WHITE,
               body_text: RGBColor = NAVY,
               col_widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = False
    if col_widths is not None:
        for col_idx, w in enumerate(col_widths):
            for cell in table.columns[col_idx].cells:
                cell.width = Inches(w)
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr_cells[i], header_fill)
        set_cell_borders(hdr_cells[i], "0B2447")
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = header_text
        run.font.size = Pt(10.5)
        run.font.name = "Calibri"
        hdr_cells[i].text = ""
        hdr_cells[i].paragraphs[0].add_run(h).font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = header_text
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10.5)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        fill = zebra[r_idx % 2]
        for c_idx, value in enumerate(row):
            shade_cell(cells[c_idx], fill)
            set_cell_borders(cells[c_idx], "B0B7C3")
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(value))
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            run.font.color.rgb = body_text
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_image(doc, name: str, *, width_in: float = 6.4, caption: str | None = None) -> None:
    path = ASSETS / name
    if not path.exists():
        body_para(doc, f"[diagram missing: {name}]", italic=True, color=RED)
        return
    doc.add_picture(str(path), width=Inches(width_in))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.font.italic = True
        r.font.size = Pt(10)
        r.font.color.rgb = GREY


def add_toc_field(doc) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-2" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:t")
    fldChar3.text = "Right-click here in Word and choose 'Update Field' to populate the table of contents."
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)
    r_element.append(fldChar4)


def add_page_break(doc) -> None:
    doc.add_page_break()
