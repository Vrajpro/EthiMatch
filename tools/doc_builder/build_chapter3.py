"""Generate Chapter 3 Methodology as Word document."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "reports" / "EthiMatch_Chapter3_Methodology.docx"

NAVY = RGBColor(0x0B, 0x24, 0x47)


def add_heading(doc, text, level=1):
    h = doc.add_heading(level=level)
    r = h.add_run(text)
    r.font.color.rgb = NAVY
    r.font.name = "Calibri"
    r.font.bold = True
    r.font.size = Pt(16) if level == 1 else Pt(13)


def add_para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.font.color.rgb = NAVY


SECTIONS = [
    ("3.1 Introduction", (
        "This chapter explains how EthiMatch was developed and evaluated. The project "
        "combines design-based software engineering with a controlled experiment that "
        "isolates whether a symbolic rule engine improves safety and explainability "
        "when the biomedical NER model is held fixed. Choices, validity threats, and "
        "constraints are discussed below."
    )),
    ("3.2 Research Approach", (
        "The study combines artefact-oriented design science with quantitative "
        "comparison. Design science fits projects whose main contribution is a "
        "technical solution evaluated against explicit requirements (Hevner et al., "
        "2004). Here the artefact is the EthiMatch pipeline and dashboard; "
        "requirements come from Chapters 1–2: efficiency, protocol compliance, "
        "explainability, and conservative handling of missing data."
        "\n\n"
        "Two decision strategies are compared on identical inputs: (A) the full "
        "neuro-symbolic pipeline; and (B) a pure-neural baseline using the same "
        "extracted entities with inclusion-oriented heuristics and no symbolic "
        "constraint checking. Holding the extractor constant attributes performance "
        "differences to the symbolic layer."
    )),
    ("3.3 System Design Methodology", (
        "System development followed an iterative, contract-driven architecture. "
        "Each pipeline stage exchanges a typed data contract: raw provider records "
        "are normalised into a PatientProfile object; neural extraction produces "
        "an ExtractedEntities structure; symbolic validation yields ValidationReport "
        "objects per trial; and the orchestrator assembles a final AuditReport for "
        "the user interface. This separation of concerns enabled independent development "
        "and testing of the data layer, neural module, rule engine, and Streamlit "
        "front end."
        "\n\n"
        "Trial eligibility criteria were externalised as JSON protocol files rather "
        "than hard-coded logic, allowing clinicians or researchers to inspect or "
        "extend protocols without modifying Python source code. A medallion-style "
        "data architecture was adopted: bronze-tier inputs (raw CSV cohorts and "
        "trial JSON), silver-tier cached neural extractions with hash-based "
        "invalidation, and gold-tier evaluation outputs. Hash-validated silver "
        "caching was introduced to mitigate neural-inference latency on CPU hardware "
        "while ensuring stale extractions are automatically refreshed when underlying "
        "notes change. Coding standards followed PEP 8 conventions throughout to "
        "maintain consistency and readability across modules."
    )),
    ("3.4 Data Sources and Cohort Construction", (
        "Two open-access datasets were used to reduce single-source bias and "
        "credentialing risk. Synthea synthetic patient CSVs (stored under "
        "data/synthea/) provide a scalable cohort with longitudinal conditions, "
        "medications, and care plans (Walonoski et al., 2018). Because Synthea does not ship clinician-written "
        "discharge summaries, the data loader synthesises EHR-style notes from "
        "structured fields, including truthful negation where comorbidities are "
        "absent—exercising the neural negation filter under controlled conditions."
        "\n\n"
        "MIMIC-IV Demo (data/mimic/) supplies a real-world structured benchmark "
        "of 100 de-identified patients with ICD-coded diagnoses and prescriptions, "
        "requiring no PhysioNet credentialing (Johnson et al., 2023). The MIMICDualSourceProvider maps "
        "these tables into the same PatientProfile contract as Synthea, enabling "
        "identical downstream processing. A procedurally generated synthetic cohort "
        "was additionally available for rapid ablation testing. Patient subsets for "
        "benchmarking were selected via configurable limits (default n = 100, "
        "maximum n = 500) to balance statistical coverage with interactive evaluation "
        "time on local hardware."
    )),
    ("3.5 Gold Standard and Ground Truth", (
        "Eligibility ground truth was derived by applying the SymbolicValidator to "
        "structured patient profiles built directly from CSV rows, independent of "
        "neural note extraction. For Synthea, only active conditions (where the STOP "
        "field is null) were included, mirroring clinical active-problem lists. "
        "This structured gold standard represents the protocol-correct label given "
        "complete structured data, and serves two purposes: evaluating the "
        "neuro-symbolic pipeline against definitive criteria, and measuring neural "
        "extraction fidelity (how closely NER-derived entities match structured "
        "fields)."
        "\n\n"
        "Verdicts were classified into three buckets for evaluation: ELIGIBLE "
        "(all mandatory inclusion criteria satisfied and no blocking exclusions), "
        "INELIGIBLE (one or more criteria definitively failed), and INCONCLUSIVE "
        "(required information missing). Treating inconclusive cases separately "
        "prevents missing data from being scored as false negatives or false "
        "positives, aligning evaluation with the system's safety-oriented design."
    )),
    ("3.6 Comparative Evaluation Design", (
        "The benchmark harness (evaluation.py) iterates over a patient cohort and "
        "six JSON-defined oncology trials. For each patient–trial pair, both decision "
        "strategies produce a binary eligibility prediction compared against the gold "
        "label. The neuro-symbolic path runs the full pipeline: silver-cache lookup, "
        "optional structured early-exit, neural extraction on cache miss, symbolic "
        "validation across all trials, and XAI report generation. The pure-neural "
        "baseline reuses the identical extracted entities but applies an "
        "inclusion-only heuristic that marks a patient eligible when extracted fields "
        "appear to satisfy inclusion criteria, without enforcing exclusions or "
        "flagging missing mandatory fields."
        "\n\n"
        "This controlled comparison directly tests the research hypothesis: if the "
        "symbolic layer contributes meaningfully, the neuro-symbolic system should "
        "exhibit lower false positive rates and improved precision relative to the "
        "baseline, with McNemar's test (McNemar, 1947) used to assess whether discordant "
        "predictions occur more often in one system's favour. Evaluation was executed through "
        "both the Streamlit dashboard and a reproducible CLI path, with results "
        "persisted to results/comparative_benchmark.json and publication figures "
        "exported to results/figures/. Figure 1 summarises this controlled design: "
        "the same patients, trials, and extractor, with only the decision layer changed."
    )),
    ("3.7 Metrics and Statistical Analysis", (
        "Performance was measured using standard information-retrieval metrics "
        "computed per dataset and macro-averaged across trials:"
        "\n\n"
        "Precision — the proportion of predicted eligible cases that are truly eligible; "
        "clinically, high precision reduces inappropriate trial offers."
        "\n"
        "Recall — the proportion of truly eligible patients correctly identified; "
        "high recall reduces missed recruitment opportunities."
        "\n"
        "F1-score — the harmonic mean of precision and recall, providing a single "
        "balanced summary."
        "\n"
        "False Positive Rate (FPR) — the proportion of ineligible patients incorrectly "
        "flagged as eligible; this was treated as the primary safety indicator."
        "\n\n"
        "Because both systems were evaluated on the same patients, McNemar's test "
        "(McNemar, 1947) was applied to the discordant prediction pairs. This procedure "
        "tests the sampling error of the difference between correlated proportions or "
        "percentages and is appropriate for paired binary classifier outcomes on "
        "identical instances, unlike independent-sample tests. A significance threshold "
        "of p < 0.05 was adopted. Discordant-cell chi-square approximation with continuity "
        "correction was implemented in evaluation.py (mcnemar_test). All experiments used "
        "fixed random seeds and content-keyed cache reads to ensure reproducibility "
        "of reported figures."
    )),
    ("3.8 Validity, Reliability and Limitations", (
        "Internal validity was strengthened by holding the neural extractor, trial "
        "registry, and patient cohort constant across compared systems. Construct "
        "validity depends on the assumption that structured CSV fields represent "
        "an appropriate gold standard; this is reasonable for Synthea and MIMIC "
        "tables but may not capture nuances present only in free-text notes. "
        "External validity is limited by the oncology-focused vocabulary, the "
        "six-trial registry, and the 100-patient MIMIC-IV Demo subset, which "
        "constrains generalisation to full hospital populations."
        "\n\n"
        "Reliability was supported through deterministic rule logic, versioned "
        "cache metadata (CACHE_VERSION and input hashes), and a smoke-test suite "
        "(scripts/qa_system_check.py) verifying registry consistency and expected "
        "verdicts on known patient profiles. Key methodological constraints include "
        "CPU-based neural inference, the use of synthesised notes for structured "
        "cohorts, and the absence of credentialed MIMIC-IV-Note free text. These "
        "constraints are acknowledged in the evaluation discussion and inform "
        "recommendations for future work."
    )),
    ("3.9 Chapter Summary", (
        "This chapter has described a design-based methodology combining contract-driven "
        "software engineering with a controlled dual-system benchmark. Dual open "
        "datasets, structured gold labels, safety-oriented metrics, and McNemar's "
        "paired testing were selected to align with the research question and the "
        "theoretical framework established in Chapter 2. The following chapter "
        "details how these methodological principles were realised in the EthiMatch "
        "artefact through its five-stage pipeline, modules, and user interface."
    )),
]


def main():
    doc = Document()
    for s in doc.sections:
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

    add_heading(doc, "Chapter 3: Research Design and Methodology", 1)

    for title, body in SECTIONS:
        add_heading(doc, title, 2)
        for para in body.split("\n\n"):
            add_para(doc, para.strip())

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))

    wc = sum(len(p.split()) for _, b in SECTIONS for p in b.split("\n\n"))
    print(f"Saved: {OUT}")
    print(f"Approximate word count: {wc}")


if __name__ == "__main__":
    main()
