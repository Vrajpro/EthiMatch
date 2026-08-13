"""Generate Chapter 2 Literature Review as Word document."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "reports" / "EthiMatch_Chapter2_Literature_Review.docx"

NAVY = RGBColor(0x0B, 0x24, 0x47)


def add_heading(doc, text, level=1):
    h = doc.add_heading(level=level)
    r = h.add_run(text)
    r.font.color.rgb = NAVY
    r.font.name = "Calibri"
    r.font.bold = True
    r.font.size = Pt(16) if level == 1 else Pt(13)


def add_para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.font.color.rgb = NAVY


SECTIONS = [
    ("2.1 Introduction", (
        "This chapter reviews literature relevant to automated clinical-trial matching "
        "in oncology. It is organised thematically around recruitment failure, "
        "biomedical text mining, neuro-symbolic reasoning, explainability, and "
        "evaluation practice, then states the research gap and theoretical framework "
        "that guide EthiMatch."
    )),
    ("2.2 Clinical Trial Recruitment and the Accrual Crisis", (
        "Patient recruitment is a leading cause of trial delay and early termination. "
        "Carlisle et al. (2015) examined 2,579 recently closed phase 2 and 3 trials "
        "and found that nearly one in five failed to achieve adequate accrual. "
        "Under-enrolment weakens statistical power and wastes participant burden. "
        "Systems that widen screening while protecting eligibility accuracy therefore "
        "matter operationally as well as academically."
        "\n\n"
        "Manual eligibility review still dominates many research settings. Coordinators "
        "must reconcile narrative notes with protocol rules on demographics, disease, "
        "biomarkers, performance status, and exclusions. Automation is best framed as "
        "a pre-screening layer that surfaces protocol-grounded evidence before human "
        "contact, not as a replacement for enrolment judgement (Zitianellis, 2025)."
    )),
    ("2.3 Biomedical Natural Language Processing and Entity Extraction", (
        "The neural component of modern clinical-matching systems depends on the "
        "ability to extract structured facts from free text. General-purpose language "
        "models perform poorly on biomedical corpora because clinical terminology, "
        "abbreviations, and negation patterns differ markedly from everyday language. "
        "Lee et al. (2020) addressed this limitation through BioBERT, a transformer "
        "model pre-trained on PubMed abstracts and MIMIC clinical notes, demonstrating "
        "superior performance on biomedical named-entity recognition (NER) and "
        "relation-extraction tasks compared with general-domain BERT (Devlin et al., 2019)."
        "\n\n"
        "Subsequent work has extended this paradigm through lighter architectures "
        "and task-specific fine-tuning. DistilBERT retains much of BERT's language "
        "understanding at lower computational cost (Sanh et al., 2019), and ClinicalBERT "
        "improves further on clinical notes (Alsentzer et al., 2019). EthiMatch therefore "
        "deploys a DistilBERT-based biomedical NER model (d4data/biomedical-ner-all, "
        "trained on MACCROBAT) via Hugging Face Transformers (Wolf et al., 2020)—a "
        "CPU-feasible compromise, not a claim that the checkpoint is BioBERT itself. "
        "Clinical negation such as \"no history of\" must not be treated as a positive "
        "finding (Chapman et al., 2001), so EthiMatch applies an explicit negation filter "
        "after NER. High NER accuracy still does not imply safe eligibility decisions: "
        "extraction errors and absent fields can propagate unless a validation layer "
        "enforces protocol rules. EthiMatch therefore treats neural extraction as a "
        "perception layer rather than a decision layer."
    )),
    ("2.4 Neuro-Symbolic Artificial Intelligence in Clinical Matching", (
        "Neuro-symbolic AI combines sub-symbolic pattern recognition with symbolic "
        "reasoning that operates over explicit rules, ontologies, or logic programs. "
        "In clinical informatics, this hybrid is attractive because medical protocols "
        "are themselves symbolic artefacts: inclusion and exclusion criteria are "
        "written as deterministic constraints that must be satisfied, not approximated. "
        "d'Avila Garcez and Lamb (2023) argue that combining neural learning with symbolic "
        "reasoning is necessary for trustworthy AI, and Sarker et al. (2022) survey "
        "current neuro-symbolic trends that EthiMatch follows: neural perception "
        "coupled to an explicit rule layer. Earlier eligibility-parsing systems such "
        "as Criteria2Query translated free-text protocol criteria into structured "
        "cohort queries (Yuan et al., 2019). Large-language-model matchers such as "
        "TrialGPT can retrieve and explain criterion-level eligibility with high "
        "accuracy, but they do not provide deterministic protocol guarantees "
        "(Jin et al., 2024). den Hamer et al. (2023) found an LLM correctly classified "
        "only about 72% of eligibility criteria, reinforcing the need for a symbolic "
        "safety layer. "
        "Loaiza-Bonilla et al. (2026) provide the most directly "
        "relevant recent evidence for oncology trial matching, reporting a prospective "
        "evaluation of a neuro-symbolic, multi-agent system integrated with an oncology "
        "knowledge graph across 3,804 patients. Their findings indicate that strict "
        "rule logic substantially improves matching safety relative to generative "
        "approaches that reason over text without protocol-level guarantees."
        "\n\n"
        "Comparing this work with purely neural pipelines highlights a recurring "
        "trade-off. Neural end-to-end models may achieve competitive recall from "
        "partial textual cues, but they are more vulnerable to false positives when "
        "negation is mishandled or required biomarkers are undocumented. Symbolic "
        "validators can refuse to guess: if a mandatory field is missing, the correct "
        "output is inconclusive. EthiMatch adopts that principle as a core safety "
        "mechanism in an open academic implementation."
    )),
    ("2.5 Explainability, Trust, and Clinical Adoption", (
        "Technical performance alone rarely determines whether clinical AI tools are "
        "adopted in practice. Lundberg and Lee (2017) introduced SHAP (SHapley "
        "Additive exPlanations), a unified framework for attributing model predictions "
        "to input features in an additive, locally faithful manner. Although SHAP was "
        "developed for general machine-learning models, its underlying idea—that "
        "decisions should be decomposed into interpretable contributions—maps naturally "
        "onto protocol-based matching, where each inclusion or exclusion criterion "
        "constitutes an auditable unit of reasoning. Complementary local explanation "
        "methods such as LIME similarly aim to make individual predictions inspectable "
        "(Ribeiro et al., 2016). The clinical cost of opaque models is well documented: "
        "Caruana et al. (2015) showed that a more accurate black-box pneumonia model "
        "hid a dangerous confounding rule that an intelligible model made visible."
        "\n\n"
        "Zitianellis (2025) makes a related point for trial pre-screening: adoption "
        "depends on transparency, clinician oversight, and conservative thresholds. "
        "EthiMatch therefore explains at two levels—feature attribution over extracted "
        "entities, and rule-level narratives stating which protocol conditions passed, "
        "failed, or could not be evaluated—without treating statistical importance as "
        "medico-legal compliance."
    )),
    ("2.6 Evaluation Datasets and Benchmarking Practice", (
        "Rigorous evaluation of clinical NLP systems requires datasets that balance "
        "accessibility, realism, and ethical feasibility. Johnson et al. (2023) "
        "describe MIMIC-IV, a large, freely accessible electronic health record "
        "repository that has become a standard benchmark for clinical machine learning. "
        "Its structured tables and (where credentialed) unstructured notes enable "
        "reproducible comparison across studies. For MSc-level research constrained "
        "by credentialing timelines, the publicly available MIMIC-IV Demo subset "
        "offers a pragmatic compromise: real clinical coding patterns and "
        "longitudinal structure without identifiable patient data."
        "\n\n"
        "Synthetic generators such as Synthea complement real EHR data by providing "
        "scalable cohorts with known ground-truth structure and no privacy restriction "
        "(Walonoski et al., 2018), which supports controlled ablation of pipeline "
        "components. A dual-dataset strategy mitigates the risk "
        "that results reflect artefacts of a single data source. Methodologically, "
        "comparative evaluation should hold the neural extractor constant when testing "
        "the contribution of symbolic validation; otherwise observed gains may be "
        "confounded by model choice rather than architecture. Paired statistical tests "
        "such as McNemar's procedure (McNemar, 1947) are appropriate when two classifiers "
        "are evaluated on identical patients, as in the EthiMatch benchmark design."
    )),
    ("2.7 Research Gap and Theoretical Framework", (
        "Synthesising the strands above reveals both convergence and a remaining gap. "
        "The literature agrees that (i) recruitment inefficiency is a major barrier "
        "to trial success; (ii) biomedical transformers improve entity extraction; "
        "(iii) neuro-symbolic designs improve safety in oncology matching; (iv) "
        "explainability and conservative handling of missing data are prerequisites "
        "for adoption; and (v) dual synthetic and real EHR benchmarks strengthen "
        "validity. What remains comparatively underdeveloped in the open academic "
        "literature is an end-to-end, reproducible prototype that integrates these "
        "elements with explicit INCONCLUSIVE semantics, JSON-governed trial "
        "protocols, paired significance testing against a pure-neural baseline, and "
        "a clinician-facing audit dashboard."
        "\n\n"
        "EthiMatch is positioned within a theoretical framework that treats "
        "clinical-trial matching as a two-stage inference problem: unstructured "
        "perception (neural NER over notes) followed by structured decision (symbolic "
        "validation over protocol rules). This framework is informed by Loaiza-Bonilla "
        "et al.'s (2026) evidence on neuro-symbolic safety, d'Avila Garcez and Lamb's (2023) "
        "account of neurosymbolic trustworthiness, Lee et al.'s (2020) domain-adapted "
        "language modelling, Lundberg and Lee's (2017) attribution principles, and "
        "Zitianellis's (2025) adoption-oriented caution regarding missing data. The "
        "hypothesis tested empirically is that separating these "
        "stages—while keeping the neural extractor fixed across compared systems—"
        "reduces false positive eligibility rates without unacceptable loss of recall."
    )),
    ("2.8 Chapter Summary", (
        "This review has traced the academic foundations of EthiMatch from the "
        "recruitment crisis through biomedical NLP, neuro-symbolic reasoning, "
        "explainability, and evaluation practice. The critical gap identified is not "
        "a lack of evidence for hybrid architectures in principle, but a shortage of "
        "transparent, reproducible implementations that clinicians and researchers "
        "can inspect, extend, and benchmark under controlled conditions. The next "
        "chapter describes how the research methodology operationalises this framework "
        "through dual-dataset experimental evaluation and paired statistical comparison."
    )),
]


def main():
    doc = Document()
    for s in doc.sections:
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

    add_heading(doc, "Chapter 2: Literature Review and Theoretical Framework", 1)

    for title, body in SECTIONS:
        add_heading(doc, title, 2)
        for para in body.split("\n\n"):
            add_para(doc, para.strip())

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))

    wc = sum(len(p.split()) for _, b in SECTIONS for p in b.split("\n\n"))
    print(f"Saved: {OUT}")
    print(f"Approximate word count: {wc}")


if __name__ == "__main__":
    main()
