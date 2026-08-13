"""Orchestrates the full EthiMatch_Design_Document.docx build.

Run from project root inside the ethimatch venv:
    .\\ethimatch\\venv\\Scripts\\python.exe build_doc.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

import build_doc_part1 as p1
import build_doc_part2 as p2
import build_doc_part3 as p3
import build_doc_part4 as p4
import build_doc_part5 as p5

PROJECT_ROOT = Path(__file__).resolve().parents[2]
OUT = PROJECT_ROOT / "docs" / "reports" / "EthiMatch_Design_Document.docx"


def _set_base_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = p1.NAVY


def _add_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = 1
    run = para.add_run("EthiMatch System Design Document  |  Vraj Dipakkumar Parekh  |  7005SCN  |  Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = p1.GREY

    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run2 = para.add_run()
    run2.font.size = Pt(8)
    run2.font.color.rgb = p1.GREY
    run2._r.append(fldChar1)
    run2._r.append(instrText)
    run2._r.append(fldChar2)


def main() -> None:
    print("[build_doc] creating document...")
    doc = Document()
    _set_base_styles(doc)

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    _add_footer(doc)

    print("  cover + toc...")
    p2.build_cover(doc)
    p2.build_toc(doc)

    print("  intro + naming + architecture...")
    p2.build_intro(doc)
    p2.build_naming(doc)
    p2.build_architecture(doc)

    print("  flowchart + use case + dataflow + components...")
    p3.build_flowchart(doc)
    p3.build_use_case(doc)
    p3.build_dataflow(doc)
    p3.build_components(doc)

    print("  file documentation A...")
    p3.build_files_intro(doc)
    p3.build_files_data(doc)

    print("  file documentation B...")
    p4.build_files_dual_providers(doc)
    p4.build_files_core(doc)
    p4.build_files_ui_eval(doc)

    print("  trials + references...")
    p4.build_trials(doc)
    p4.build_references_mapping(doc)

    print("  evaluation + rationale + pros/cons + risks + conclusion...")
    p5.build_evaluation(doc)
    p5.build_rationale(doc)
    p5.build_pros_cons(doc)
    p5.build_risks(doc)
    p5.build_future_glossary_conclusion(doc)

    doc.save(str(OUT))
    print(f"[build_doc] SAVED: {OUT}")


if __name__ == "__main__":
    main()
