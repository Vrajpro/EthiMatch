"""Shared Word styling helpers for EthiMatch project report."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[2]
FIGURES = PROJECT_ROOT / "docs" / "figures"

NAVY = RGBColor(0x0B, 0x24, 0x47)
TEAL = RGBColor(0x19, 0xA7, 0xCE)
GREY = RGBColor(0x4A, 0x55, 0x68)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def setup_document() -> Document:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(12)
    normal.font.color.rgb = NAVY
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    return doc


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def add_chapter_heading(doc: Document, text: str) -> None:
    h = doc.add_heading(level=1)
    r = h.add_run(text)
    r.font.name = "Calibri"
    r.font.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = NAVY


def add_section_heading(doc: Document, text: str) -> None:
    h = doc.add_heading(level=2)
    r = h.add_run(text)
    r.font.name = "Calibri"
    r.font.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = NAVY


def add_para(doc: Document, text: str, *, italic: bool = False, center: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.font.color.rgb = NAVY
    r.font.italic = italic


def add_bullet(doc: Document, text: str) -> None:
    if not text.strip():
        return
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    caption: str | None = None,
) -> None:
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cap.paragraph_format.space_before = Pt(8)
        cap.paragraph_format.space_after = Pt(4)
        r = cap.add_run(caption)
        r.font.italic = True
        r.font.size = Pt(10)
        r.font.color.rgb = GREY
        r.font.name = "Calibri"
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, "0B2447")
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = "Calibri"
    for r_idx, row in enumerate(rows):
        fill = "F5F7FA" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            shade_cell(cell, fill)
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(value))
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            run.font.color.rgb = NAVY
    doc.add_paragraph()


def add_image(doc: Document, filename: str, caption: str, width: float = 5.8) -> None:
    path = FIGURES / filename
    if not path.exists():
        add_para(doc, f"[Figure not found: {filename}]", italic=True)
        return
    # UI screenshots are taller; keep them slightly narrower so captions stay readable
    if "screenshots/" in filename.replace("\\", "/"):
        width = min(width, 5.4)
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = GREY


def add_toc_field(doc: Document) -> None:
    add_para(
        doc,
        "Right-click below and choose Update Field → Update entire table "
        "to refresh page numbers in Microsoft Word.",
        italic=True,
    )
    p = doc.add_paragraph()
    run = p.add_run()
    for tag, text in (
        ("begin", None),
        ("instr", 'TOC \\o "1-2" \\h \\z \\u'),
        ("separate", "Table of Contents"),
        ("end", None),
    ):
        if tag == "instr":
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = text
        elif tag == "separate":
            el = OxmlElement("w:t")
            el.text = text
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), "separate")
            run._r.append(fld)
            run._r.append(el)
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), "end")
            run._r.append(fld)
            continue
        else:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), tag)
        run._r.append(el)


def render_sections(doc: Document, sections: list[tuple[str, str]]) -> None:
    for title, body in sections:
        add_section_heading(doc, title)
        for para in body.split("\n\n"):
            text = para.strip()
            if not text:
                continue
            if text.startswith("• "):
                for line in text.split("\n"):
                    add_bullet(doc, line.lstrip("• ").strip())
            else:
                add_para(doc, text)


def word_count(sections: list[tuple[str, str]]) -> int:
    return sum(len(p.split()) for _, b in sections for p in b.split("\n\n") if p.strip())
