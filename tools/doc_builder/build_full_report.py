"""Build the complete EthiMatch 7005SCN Project Report (all chapters) as Word.

Run from project root:
    .\\ethimatch\\venv\\Scripts\\python.exe tools\\doc_builder\\build_full_report.py

Output:
    docs/reports/EthiMatch_Project_Report.docx
"""
from __future__ import annotations

import sys
from pathlib import Path

# Allow imports from this package directory
sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from report_all_chapters import (
    ABSTRACT,
    ALL_CHAPTERS,
    CHAPTER_6_TABLE,
    REFERENCES,
)
from report_styles import (
    FIGURES,
    GREY,
    NAVY,
    TEAL,
    add_chapter_heading,
    add_image,
    add_page_break,
    add_para,
    add_section_heading,
    add_table,
    add_toc_field,
    render_sections,
    setup_document,
    word_count,
)

PROJECT_ROOT = Path(__file__).resolve().parents[2]
OUT = PROJECT_ROOT / "docs" / "reports" / "EthiMatch_Project_Report.docx"


def build_cover(doc) -> None:
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Coventry University")
    r.font.size = Pt(14)
    r.font.color.rgb = GREY
    r.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("7005SCN Individual Research Project")
    r.font.size = Pt(13)
    r.font.color.rgb = GREY
    r.font.name = "Calibri"

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("EthiMatch")
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = NAVY
    r.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("A Neuro-Symbolic AI System for\nClinical-Trial Matching")
    r.font.size = Pt(16)
    r.font.color.rgb = TEAL
    r.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Source code: ")
    r.font.size = Pt(12)
    r.font.color.rgb = NAVY
    r.font.name = "Calibri"
    r2 = p.add_run("https://github.com/Vrajpro/EthiMatch")
    r2.font.size = Pt(12)
    r2.font.color.rgb = TEAL
    r2.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(
        "Examiner Q&A (if any questions arise while reading this PDF): "
    )
    r.font.size = Pt(11)
    r.font.color.rgb = NAVY
    r.font.name = "Calibri"
    r2 = p.add_run(
        "https://github.com/Vrajpro/EthiMatch/blob/main/docs/EXAMINER_FAQ.md"
    )
    r2.font.size = Pt(11)
    r2.font.color.rgb = TEAL
    r2.font.name = "Calibri"

    for _ in range(2):
        doc.add_paragraph()

    add_table(
        doc,
        ["Field", "Detail"],
        [
            ["Assignment", "CW2 — Project Report"],
            ["Student", "Vraj Dipakkumar Parekh"],
            ["Student ID", "16485659"],
            ["Course", "MSc Data Science"],
            ["Supervisor", "Someyah Bazin"],
            ["Submission", "August 2026"],
            ["Word count (body)", "approximately 7,000 (±10%)"],
            ["Submission format", "PDF only (Turnitin)"],
        ],
    )
    add_page_break(doc)


def build_abstract(doc) -> None:
    add_chapter_heading(doc, "Abstract")
    add_para(doc, ABSTRACT)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("Keywords: ")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.font.color.rgb = NAVY
    r2 = p.add_run(
        "clinical trials; neuro-symbolic AI; biomedical NLP; patient matching; "
        "explainable AI; oncology"
    )
    r2.font.name = "Calibri"
    r2.font.size = Pt(12)
    r2.font.color.rgb = NAVY
    add_page_break(doc)


def build_toc(doc) -> None:
    add_chapter_heading(doc, "Table of Contents")
    add_toc_field(doc)
    add_page_break(doc)
    add_chapter_heading(doc, "List of Figures")
    add_para(doc, "Figure 1 — Controlled evaluation design (neuro-symbolic vs pure neural).")
    add_para(doc, "Figure 2 — EthiMatch primary stakeholder use cases.")
    add_para(doc, "Figure 3 — EthiMatch five-stage neuro-symbolic architecture.")
    add_para(doc, "Figure 4 — End-to-end pipeline control flow (cache, early-exit, NER, rules, XAI).")
    add_para(doc, "Figure 5 — Data flow from sources through pipeline to AuditReport.")
    add_para(doc, "Figure 6 — EthiMatch Dashboard: registered trials, Synthea cohort overview, and session controls.")
    add_para(doc, "Figure 7 — Patient Matching: Quick Entry result showing Eligible verdict with expanded patient detail and audit tabs.")
    add_para(doc, "Figure 8 — Cohort Discovery: trial-centric screening of 100 patients against ONC-001 (eligible, inconclusive, and ineligible counts).")
    add_para(doc, "Figure 9 — Evaluation page: neuro-symbolic vs pure-neural benchmark settings with cached synthetic results and McNemar summary.")
    add_para(doc, "Figure 10 — Logical component map (data, neural, symbolic, XAI, UI).")
    add_para(doc, "Figure 11 — EthiMatch Evaluation page 2D comparison chart (synthetic, n = 100).")
    add_page_break(doc)
    add_chapter_heading(doc, "List of Tables")
    add_para(doc, "Table 1 — Synthetic comparative metrics (n = 100, 6 trials).")
    add_para(doc, "Table 2 — McNemar paired comparison (synthetic).")
    add_para(doc, "Table 3 — Cross-source F1 and FPR comparison.")
    add_para(doc, "Table 4 — ONC-001 full Synthea safety comparison.")
    add_para(doc, "Table 5 — Explicit claim discipline for evaluation findings.")
    add_para(doc, "Table 6 — Objective–outcome closure for examiners.")
    add_para(doc, "Table 7 — Final claim discipline summary.")
    add_page_break(doc)


def build_chapter(doc, title: str, sections, *, figures: list[tuple[str, str]] | None = None) -> int:
    add_chapter_heading(doc, title)
    fig_map = {key: cap for key, cap in (figures or [])}
    for sec_title, body in sections:
        add_section_heading(doc, sec_title)
        for para in body.split("\n\n"):
            text = para.strip()
            if not text:
                continue
            if text.startswith("•"):
                from report_styles import add_bullet
                for line in text.split("\n"):
                    add_bullet(doc, line.lstrip("• ").strip())
            else:
                add_para(doc, text)
        if sec_title in fig_map:
            add_image(doc, fig_map[sec_title][0], fig_map[sec_title][1])
    return word_count(sections)


def build_chapter6_extras(doc) -> None:
    add_section_heading(doc, "Table 1 — Comparative benchmark summary")
    add_table(
        doc,
        [
            "Dataset",
            "Neuro P", "Neuro R", "Neuro F1", "Neuro FPR",
            "Neural P", "Neural R", "Neural F1", "Neural FPR",
            "McNemar p",
        ],
        CHAPTER_6_TABLE,
    )
    add_image(
        doc,
        "07_evaluation.png",
        "Figure 1 — Comparative evaluation methodology (neuro-symbolic vs pure-neural).",
    )
    add_image(
        doc,
        "01_architecture.png",
        "Figure 3 — Five-stage pipeline evaluated in benchmarks.",
        width=5.5,
    )


def build_references(doc) -> None:
    add_chapter_heading(doc, "References")
    for ref in REFERENCES:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(ref)
        r.font.name = "Calibri"
        r.font.size = Pt(11)
        r.font.color.rgb = NAVY


def build_ai_declaration(doc) -> None:
    add_page_break(doc)
    add_chapter_heading(doc, "AI Usage Declaration")
    add_para(
        doc,
        "I acknowledge the use of AI to generate materials that were included within my "
        "final submission.",
    )
    add_para(
        doc,
        "This assessment is in the amber category. AI was used only for assistance with "
        "comprehension, spelling, language, grammar, and software-development support. "
        "It was not used to invent results, citations, or evaluation numbers.",
    )
    add_para(
        doc,
        "Tools used: Cursor (Composer) for code structure, report drafting support, and "
        "diagram generation; Grammarly for spelling and grammar. Dates of access: June 2026 "
        "to August 2026.",
    )
    add_para(
        doc,
        "How material was generated: prompts asked the tool to draft chapter outlines, "
        "rephrase technical descriptions of the EthiMatch pipeline, generate architecture "
        "diagrams from the codebase, and check consistency of evaluation numbers against "
        "results/comparative_benchmark.json.",
    )
    add_para(
        doc,
        "What the output was: draft Word chapter text, Python diagram scripts, and "
        "suggested wording for methods, evaluation, and conclusion sections.",
    )
    add_para(
        doc,
        "How the output was changed by the author: all evaluation figures were replaced "
        "with values from the saved benchmark files (synthetic n = 100; F1 65.5% vs 56.2%; "
        "FPR 0.6% vs 2.4%; McNemar p ≈ 0.067, not significant at 0.05). Over-claims were "
        "removed. References were checked (including McNemar, 1947). Chapter 1 aims, "
        "Chapter 5 timeline, and the claim-discipline tables were edited so they match "
        "the implemented artefact. The author remains responsible for the submitted argument.",
    )


def add_footer(doc) -> None:
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(
        "EthiMatch Project Report | Vraj Dipakkumar Parekh | 7005SCN | Page "
    )
    run.font.size = Pt(8)
    run.font.color.rgb = GREY
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    r2 = para.add_run()
    r2.font.size = Pt(8)
    r2.font.color.rgb = GREY
    r2._r.append(fld1)
    r2._r.append(instr)
    r2._r.append(fld2)


def main() -> None:
    print("[full_report] building complete project report...")
    if not FIGURES.exists():
        print("[full_report] WARNING: docs/figures/ missing — run build_diagrams.py first")

    doc = setup_document()
    add_footer(doc)

    build_cover(doc)
    build_abstract(doc)
    build_toc(doc)

    total_words = len(ABSTRACT.split())

    chapter_figures = {
        "Chapter 3: Research Design and Methodology": {
            "3.6 Comparative Evaluation Design": (
                "07_evaluation.png",
                "Figure 1 — Controlled evaluation design.",
            ),
        },
        "Chapter 4: Artefact Design and Development": {
            "4.3 Five-Stage Pipeline Architecture": (
                "01_architecture.png",
                "Figure 3 — EthiMatch five-stage neuro-symbolic pipeline.",
            ),
            "4.4 Data Layer": (
                "02_dataflow.png",
                "Figure 5 — Data flow from inputs to audit report.",
            ),
        },
    }

    for title, sections in ALL_CHAPTERS:
        figs = chapter_figures.get(title, {})
        # Convert to list of tuples for render - we'll handle in custom loop
        add_chapter_heading(doc, title)
        wc = 0
        for sec_title, body in sections:
            add_section_heading(doc, sec_title)
            for para in body.split("\n\n"):
                text = para.strip()
                if not text:
                    continue
                if text.lstrip().startswith("•"):
                    from report_styles import add_bullet
                    for line in text.split("\n"):
                        cleaned = line.strip().lstrip("•").strip()
                        add_bullet(doc, cleaned)
                else:
                    add_para(doc, text)
                wc += len(text.split())
            if sec_title in figs:
                fname, cap = figs[sec_title]
                add_image(doc, fname, cap)

        if title == "Chapter 6: Evaluation and Discussion":
            build_chapter6_extras(doc)

        total_words += wc
        print(f"  {title}: ~{wc} words")
        add_page_break(doc)

    build_references(doc)
    build_ai_declaration(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"[full_report] SAVED: {OUT}")
    print(f"[full_report] Approximate body word count: {total_words}")
    print("[full_report] Open in Word, update Table of Contents, then Save as PDF for Turnitin")


if __name__ == "__main__":
    main()
